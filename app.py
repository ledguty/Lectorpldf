import streamlit as st
import fitz
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def es_cursiva(font):
    return "italic" in font.lower() or "oblique" in font.lower()

def modificar_pdf(archivo):
    try:
        doc = fitz.open(stream=archivo.read(), filetype="pdf")
        contenido = []
        
        for pagina in doc:
            bloques = pagina.get_text("dict")["blocks"]
            for bloque in bloques:
                if "lines" in bloque:
                    parrafo = ""
                    for linea in bloque["lines"]:
                        for span in linea["spans"]:
                            texto = span["text"]
                            if es_cursiva(span["font"]):
                                texto = f"_{texto}_"
                            parrafo += texto
                        parrafo += "\n"
                    contenido.append({
                        "texto": parrafo.strip(),
                        "bbox": bloque["bbox"],
                        "tipo": "texto"
                    })
                elif "image" in bloque:
                    contenido.append({
                        "bbox": bloque["bbox"],
                        "tipo": "imagen"
                    })
        
        doc.close()
        return contenido
    except Exception as e:
        st.error(f"Error al procesar el archivo: {str(e)}")
        return None

def procesar_texto(texto):
    lineas = texto.split('\n')
    texto_procesado = ""
    for i in range(len(lineas)):
        linea_actual = lineas[i].strip()
        if linea_actual:
            if i < len(lineas) - 1 and not linea_actual.endswith(('.', ':', ')', '?', '_')) and not linea_actual[-1].isdigit():
                texto_procesado += linea_actual + " "
            else:
                texto_procesado += linea_actual + "\n"
    return texto_procesado

def crear_documento_word(contenido):
    doc = Document()
    for item in contenido:
        if item["tipo"] == "texto":
            texto_procesado = procesar_texto(item["texto"])
            p = doc.add_paragraph()
            partes = re.split(r'(_[^_]+_)', texto_procesado)
            for parte in partes:
                if parte.startswith('_') and parte.endswith('_'):
                    run = p.add_run(parte[1:-1])
                    run.italic = True
                else:
                    p.add_run(parte)
            
            # Alineación basada en la posición del bloque
            if item["bbox"][0] < 100:  # Ajusta este valor según sea necesario
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            elif item["bbox"][0] > 300:  # Ajusta este valor según sea necesario
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        elif item["tipo"] == "imagen":
            # Aquí podrías añadir un marcador de posición para las imágenes
            doc.add_paragraph("[Imagen]")
    
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    return output_buffer.getvalue()

st.title("Convertir PDF a Word con formato preservado")
st.write("Sube un archivo PDF para convertirlo a Word, preservando el formato y aplicando reglas específicas de salto de línea.")

archivo = st.file_uploader("Subir archivo PDF", type="pdf")

if archivo is not None:
    if st.button("Procesar PDF"):
        contenido_modificado = modificar_pdf(archivo)
        if contenido_modificado:
            docx_modificado = crear_documento_word(contenido_modificado)
            st.download_button(
                label="Descargar documento Word modificado",
                data=docx_modificado,
                file_name="documento_modificado.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
